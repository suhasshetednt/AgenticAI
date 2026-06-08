"""ASL Airlines brand styling helpers for python-docx documents."""

from __future__ import annotations

import logging
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

COLOR_NAVY = RGBColor(0x00, 0x38, 0x6B)
COLOR_GOLD = RGBColor(0xC8, 0x9A, 0x00)
COLOR_DARK = RGBColor(0x1A, 0x1A, 0x1A)
# OCC / Eagle Eye design tokens
_OCC_TEAL = RGBColor(0x0F, 0x47, 0x61)       # heading colour
_HEADER_BG = "95B3D7"                          # light-blue content table header
_HEADER_BG_DARK = "595959"                     # dark-grey meta / cover table header
_BORDER_COLOR = "4F81BD"                       # OCC blue border


def _set_cell_bg(cell: Any, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _try_paragraph_style(doc: Any, *style_names: str) -> tuple[Any, str]:
    """Try style names in order; return (paragraph, resolved_style_name).

    Falls back to a plain unstyled paragraph if none of the named styles exist
    in the document. Never raises KeyError.
    """
    for style in style_names:
        try:
            para = doc.add_paragraph(style=style)
            return para, style
        except KeyError:
            logger.debug("Style %r not found in document, trying next", style)
    # Last-resort: no style at all
    para = doc.add_paragraph()
    return para, ""


def add_heading(doc: Any, text: str, level: int) -> None:
    try:
        para = doc.add_heading(text, level=level)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # OCC template defines teal (0F4761) via styles — don't override.
    except (KeyError, ValueError):
        logger.debug("Heading %d style not found, falling back to bold paragraph", level)
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(max(20 - level * 2, 10))
        run.font.color.rgb = _OCC_TEAL


def add_paragraph(doc: Any, text: str, size: int = 10) -> None:
    para = doc.add_paragraph(text)
    if para.runs:
        para.runs[0].font.size = Pt(size)


def add_bullet(doc: Any, text: str) -> None:
    """Add a bullet paragraph, gracefully handling missing styles.

    Priority order:
      1. "List Bullet"      — standard python-docx built-in
      2. "List Paragraph"   — used by some ASL branded templates
      3. "Normal"           — guaranteed to exist in any .docx
      4. unstyled paragraph — absolute last resort
    """
    para, resolved = _try_paragraph_style(
        doc, "List Bullet", "List Paragraph", "Normal"
    )

    # "List Bullet" provides its own bullet glyph via numbering XML.
    # Every other style needs a manual bullet character prepended.
    bullet_text = text if resolved == "List Bullet" else f"\u2022 {text}"
    run = para.add_run(bullet_text)
    run.font.size = Pt(10)


def add_code(doc: Any, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def _apply_occ_borders(table: Any) -> None:
    """Apply OCC-style blue (4F81BD) borders to all table edges."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), _BORDER_COLOR)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_table(
    doc: Any,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = _HEADER_BG,
) -> None:
    """Add an OCC-styled table with coloured header row and blue borders.

    Default header fill is light-blue (95B3D7, black text).
    Pass header_fill=_HEADER_BG_DARK for dark-grey (595959, white text) meta tables.
    """
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)

    dark_fill = header_fill.upper() in ("595959", "1F3864", "00386B")

    # Header row
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _set_cell_bg(cell, header_fill)
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if dark_fill else RGBColor(0x00, 0x00, 0x00)

    # Data rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(cols):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    _apply_occ_borders(table)
    doc.add_paragraph()


def set_page_margins(doc: Any) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)