"""DocxRenderer — parse markdown (markdown-it-py) and render a branded .docx."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from markdown_it import MarkdownIt
from markdown_it.token import Token

from adl_automated_delivery_pipeline.documentation import brand
from adl_automated_delivery_pipeline.documentation.context import DocContext

logger = logging.getLogger(__name__)

_MD = MarkdownIt("commonmark").enable("table")


def _inline_text(token: Token) -> str:
    if token.children:
        return "".join(child.content for child in token.children)
    return token.content


def _parse_table(tokens: list[Token], start: int) -> tuple[list[str], list[list[str]], int]:
    headers: list[str] = []
    rows: list[list[str]] = []
    current: list[str] = []
    in_body = False
    j = start + 1
    while tokens[j].type != "table_close":
        ttype = tokens[j].type
        if ttype == "thead_open":
            in_body = False
        elif ttype == "tbody_open":
            in_body = True
        elif ttype == "tr_open":
            current = []
        elif ttype == "tr_close" and in_body:
            rows.append(current)
        elif ttype in ("th_open", "td_open"):
            text = _inline_text(tokens[j + 1])
            (headers if ttype == "th_open" else current).append(text)
            j += 2  # advance past the inline child; the loop's j += 1 then skips the close tag
        j += 1
    return headers, rows, j + 1


def _render_tokens(doc: Any, tokens: list[Token]) -> None:
    i = 0
    n = len(tokens)
    while i < n:
        tok = tokens[i]
        ttype = tok.type
        if ttype == "heading_open":
            brand.add_heading(doc, _inline_text(tokens[i + 1]), level=int(tok.tag[1]))
            i += 3
        elif ttype == "paragraph_open":
            brand.add_paragraph(doc, _inline_text(tokens[i + 1]))
            i += 3
        elif ttype == "bullet_list_open":
            # Track nesting depth so a nested sub-list's close does not end the
            # outer list early. Nested items are flattened to bullets.
            depth = 1
            i += 1
            while depth > 0:
                inner = tokens[i].type
                if inner == "bullet_list_open":
                    depth += 1
                elif inner == "bullet_list_close":
                    depth -= 1
                elif inner == "inline":
                    brand.add_bullet(doc, _inline_text(tokens[i]))
                i += 1
        elif ttype == "table_open":
            headers, rows, i = _parse_table(tokens, i)
            brand.add_table(doc, headers, rows)
        elif ttype in ("fence", "code_block"):
            brand.add_code(doc, tok.content.rstrip("\n"))
            i += 1
        else:
            i += 1


class DocxRenderer:
    extension = "docx"

    def render(self, markdown: str, out_path: Path, context: DocContext) -> Path:
        # `context` is part of the Renderer protocol and reserved for future metadata use.
        out_path = out_path.with_suffix(".docx")
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        brand.set_page_margins(doc)
        _render_tokens(doc, _MD.parse(markdown))
        doc.save(str(out_path))
        return out_path
