"""
Documentation Agent — ASL Airlines / ADL Project
==================================================
Generates a professional Technical Implementation Document (.docx) from
structured Jira ticket requirements and an optional SQL query / VDS path.

Uses Claude (single structured JSON call) to populate every section, then
assembles the Word document with python-docx applying proper heading levels,
bullet lists, and compact tables.

Output directory: <cwd>/Project Documentation/
Filename pattern: {ticket_id}_{YYYYMMDD_HHMMSS}.docx
"""
from __future__ import annotations

import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, cast

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from langchain_core.messages import HumanMessage, SystemMessage

from adl_automated_delivery_pipeline.config import settings
from adl_automated_delivery_pipeline.llm import make_claude

# Importing ``settings`` triggers env-file discovery/loading once, so this
# module needs no separate dotenv bootstrap.

logger = logging.getLogger(__name__)

_OUTPUT_DIR = Path.cwd() / "Project Documentation"

# ASL brand colours
_COLOR_NAVY = RGBColor(0x00, 0x38, 0x6B)   # ASL navy blue
_COLOR_GOLD = RGBColor(0xC8, 0x9A, 0x00)   # accent gold
_COLOR_DARK = RGBColor(0x1A, 0x1A, 0x1A)   # near-black text
_COLOR_HEADER_BG = "00386B"                 # table header fill (hex string for XML)


# ── LLM prompt ──────────────────────────────────────────────────────────────────

def _build_doc_prompt(
    reqs: Any,
    sql: str = "",
    vds_path: str = "",
) -> str:
    """Build the single LLM prompt that requests all document sections as JSON.

    Args:
        reqs: A TicketRequirements dataclass instance.
        sql: Optional generated SQL query string.
        vds_path: Optional Dremio VDS output path.

    Returns:
        Formatted prompt string for the LLM.
    """
    fields_txt = "\n".join(
        f"  - {f['name']}: {f.get('description', '')}" for f in reqs.output_fields
    ) or "  - (not specified)"
    tables_txt = "\n".join(f"  - {t}" for t in reqs.source_tables) or "  - (not specified)"
    transforms_txt = "\n".join(f"  - {t}" for t in reqs.transformations) or "  - (not specified)"
    filters_txt = "\n".join(f"  - {c}" for c in reqs.filter_conditions) or "  - None specified"
    criteria_txt = "\n".join(f"  - {a}" for a in reqs.acceptance_criteria) or "  - (not specified)"
    sql_block = f"\n--- GENERATED SQL ---\n{sql}" if sql.strip() else ""
    vds_block = f"\n--- VDS OUTPUT PATH ---\n{vds_path}" if vds_path.strip() else ""
    notes_block = f"\n--- EXTRA NOTES ---\n{reqs.extra_notes}" if reqs.extra_notes else ""

    return f"""You are a technical documentation writer for ASL Airlines DataLake team (DnT Infotech).
Generate a complete Technical Implementation Document for a Dremio/AMOS data integration ticket.

Return ONLY a single valid JSON object — no markdown fences, no preamble, no explanation.
All string values must be plain text (no markdown inside the JSON strings).

--- TICKET INFORMATION ---
Ticket ID  : {reqs.ticket_id}
Summary    : {reqs.summary}

--- BUSINESS REQUIREMENT ---
{reqs.business_requirement}

--- SOURCE DATABASE ---
{reqs.source_database}

--- SOURCE TABLES ---
{tables_txt}

--- OUTPUT FIELDS ---
{fields_txt}

--- TRANSFORMATION RULES ---
{transforms_txt}

--- FILTER CONDITIONS ---
{filters_txt}

--- ACCEPTANCE CRITERIA ---
{criteria_txt}{sql_block}{vds_path}{vds_block}{notes_block}

Generate the JSON document with EXACTLY this structure:
{{
  "project": "ASL Airlines — ADL DataLake Integration",
  "team": "DnT Infotech - DataLake Team",
  "stakeholders": [
    "Tony Fourmeau — Technical Lead, ASL Airlines",
    "Suhas Shete — Development Lead, DnT Infotech"
  ],
  "initial_requirements_summary": "<Two to three sentence narrative paragraph explaining the business context and what this implementation delivers>",
  "initial_requirements_items": [
    "<VDS 1: describe what this VDS provides>",
    "<VDS 2: if applicable>",
    "<Key output or report this enables>"
  ],
  "additional_requirements": [
    "<Edge case or business logic rule 1>",
    "<Null/exclusion handling rule>",
    "<Data quality or validation requirement>",
    "<Performance or scope constraint>"
  ],
  "risks": [
    {{"risk": "<Technical or data risk>", "mitigation": "<Concrete mitigation step>"}},
    {{"risk": "<Schema or data type risk>", "mitigation": "<Mitigation>"}},
    {{"risk": "<Performance risk for large tables>", "mitigation": "<Mitigation>"}}
  ],
  "process_continuity": [
    {{"issue": "<What happens if source data is missing>", "mitigation": "<Fallback behavior>"}},
    {{"issue": "<What happens if the VDS is unavailable>", "mitigation": "<Recovery step>"}}
  ],
  "data_sources": [
    {{"name": "<table_name>", "path": "<full_source_path e.g. amos_postgres.amos.table_name>", "description": "<what this table contains>"}}
  ],
  "interface_process": [
    {{
      "component": "<VDS or layer name>",
      "steps": [
        "<Step 1: what data is selected and from where>",
        "<Step 2: joins or transformations applied>",
        "<Step 3: filters and business logic>",
        "<Step 4: output fields and naming conventions>"
      ]
    }}
  ],
  "data_dictionary": [
    {{"field": "<output_field_name>", "source_field": "<source_table.column_name>", "logic": "<how this field is derived — transformation, cast, CASE WHEN, etc.>"}}
  ],
  "milestones": [
    {{"milestone": "Requirements review and sign-off", "days": 1}},
    {{"milestone": "VDS development and unit testing", "days": 3}},
    {{"milestone": "Data validation against source", "days": 2}},
    {{"milestone": "UAT and stakeholder review", "days": 2}},
    {{"milestone": "Production deployment", "days": 1}}
  ],
  "vds_source_tables": [
    {{"name": "<table_name>", "path": "<amos_postgres.amos.table_name>"}}
  ],
  "vds_output_paths": [
    {{"name": "<vds_name>", "path": "<dremio-db.folder.vds_name>"}}
  ],
  "key_business_rules": [
    "<Rule 1 — e.g. date range or status filter>",
    "<Rule 2 — null handling>",
    "<Rule 3 — deduplication or grouping logic>"
  ],
  "access_control": "Developer and Tester roles have SELECT access to the VDS output paths. Write access is restricted to the DataLake admin role.",
  "signoff": [
    {{"role": "Technical Lead", "name": "Tony Fourmeau (ASL)", "date": ""}},
    {{"role": "Development Lead", "name": "Suhas Shete (DnT)", "date": ""}},
    {{"role": "QA / Tester", "name": "", "date": ""}},
    {{"role": "Business Owner", "name": "", "date": ""}}
  ]
}}

Important:
- Populate every field meaningfully based on the ticket information above.
- data_dictionary must have one row per output field listed above.
- vds_output_paths must reflect the actual VDS path if provided, or a reasonable default under dremio-db.
- interface_process steps must describe the actual SQL logic (joins, filters, field derivations).
- milestones days must be realistic integers summing to a total of 9 to 14 working days.
"""


# ── Document builder helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell: Any, hex_color: str) -> None:
    """Set table cell background fill colour using OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_cell(cell: Any, text: str, font_size: int = 10, color: RGBColor | None = None) -> None:
    """Write bold text into a table cell."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Any, text: str, level: int) -> None:
    """Add a styled heading paragraph."""
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.color.rgb = _COLOR_NAVY if level == 1 else _COLOR_DARK
        run.font.bold = True


def _add_bullet(doc: Any, text: str, level: int = 0) -> None:
    """Add a list bullet paragraph."""
    para = doc.add_paragraph(text, style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in para.runs:
        run.font.size = Pt(10)


def _add_table_with_header(
    doc: Any,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """Add a compact, styled table with a navy header row.

    Args:
        doc: The active Document object.
        headers: Column header labels.
        rows: List of row data; each row is a list of strings aligned to headers.
    """
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        _set_cell_bg(hdr_cells[idx], _COLOR_HEADER_BG)
        _bold_cell(hdr_cells[idx], header_text, font_size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        data_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            data_cells[col_idx].text = str(cell_text)
            for para in data_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer after table


def _add_key_value(doc: Any, label: str, value: str) -> None:
    """Add a bold-label followed by plain value on the same paragraph."""
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_value = para.add_run(value)
    run_value.font.size = Pt(10)


# ── Core document assembly ───────────────────────────────────────────────────────

def _assemble_document(
    reqs: Any,
    content: dict[str, Any],
    sql: str,
    vds_path: str,
) -> Any:
    """Build the python-docx Document from LLM-generated content.

    Args:
        reqs: The TicketRequirements dataclass.
        content: Parsed JSON dict from LLM response.
        sql: The generated SQL string (may be empty).
        vds_path: The VDS output path (may be empty).

    Returns:
        A fully populated Document object ready to save.
    """
    doc = Document()

    # ── Page margins (narrower for denser layout) ──────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Title block ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("ASL AIRLINES — TECHNICAL IMPLEMENTATION DOCUMENT")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = _COLOR_NAVY
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(f"{reqs.ticket_id}  |  {reqs.summary}")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = _COLOR_GOLD
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_str = datetime.now(timezone.utc).strftime("%d %B %Y")
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Prepared: {date_str}  |  DnT Infotech - DataLake Team")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── 1. Project ─────────────────────────────────────────────────────────────
    _add_heading(doc, "1. Project", level=1)
    _add_key_value(doc, "Project", content.get("project", "ASL Airlines — ADL DataLake Integration"))
    doc.add_paragraph()

    # ── 2. Team ───────────────────────────────────────────────────────────────
    _add_heading(doc, "2. Team", level=1)
    _add_key_value(doc, "Team", content.get("team", "DnT Infotech - DataLake Team"))
    doc.add_paragraph()

    # ── 3. Stakeholders ───────────────────────────────────────────────────────
    _add_heading(doc, "3. Stakeholders", level=1)
    for stakeholder in content.get("stakeholders", []):
        _add_bullet(doc, stakeholder)
    doc.add_paragraph()

    # ── 4. Initial Requirements ───────────────────────────────────────────────
    _add_heading(doc, "4. Initial Requirements", level=1)
    summary_text = content.get("initial_requirements_summary", reqs.business_requirement)
    doc.add_paragraph(summary_text).runs[0].font.size = Pt(10)
    doc.add_paragraph()
    for item in content.get("initial_requirements_items", []):
        _add_bullet(doc, item)
    doc.add_paragraph()

    # ── 5. Additional Requirements ────────────────────────────────────────────
    _add_heading(doc, "5. Additional Requirements", level=1)
    for item in content.get("additional_requirements", []):
        _add_bullet(doc, item)
    doc.add_paragraph()

    # ── 6. Risks & Mitigation ─────────────────────────────────────────────────
    _add_heading(doc, "6. Risks & Mitigation", level=1)
    risks = content.get("risks", [])
    if risks:
        _add_table_with_header(
            doc,
            headers=["Risk", "Mitigation"],
            rows=[[r.get("risk", ""), r.get("mitigation", "")] for r in risks],
        )
    else:
        doc.add_paragraph("No significant risks identified.")
        doc.add_paragraph()

    # ── 7. Process Continuity ─────────────────────────────────────────────────
    _add_heading(doc, "7. Process Continuity", level=1)
    continuity = content.get("process_continuity", [])
    if continuity:
        _add_table_with_header(
            doc,
            headers=["Issue", "Mitigation"],
            rows=[[c.get("issue", ""), c.get("mitigation", "")] for c in continuity],
        )
    else:
        doc.add_paragraph("No process continuity concerns identified.")
        doc.add_paragraph()

    # ── 8. Implementation ─────────────────────────────────────────────────────
    _add_heading(doc, "8. Implementation", level=1)

    # 8a. Data Sources
    _add_heading(doc, "8.1  Data Sources", level=2)
    data_sources = content.get("data_sources", [])
    if data_sources:
        _add_table_with_header(
            doc,
            headers=["Table Name", "Full Path", "Description"],
            rows=[
                [
                    ds.get("name", ""),
                    ds.get("path", ""),
                    ds.get("description", ""),
                ]
                for ds in data_sources
            ],
        )
    else:
        for tbl in reqs.source_tables:
            _add_bullet(doc, tbl)
        doc.add_paragraph()

    # 8b. Interface Process
    _add_heading(doc, "8.2  Interface Process", level=2)
    interface_process = content.get("interface_process", [])
    if interface_process:
        for component in interface_process:
            component_name = component.get("component", "")
            _add_heading(doc, component_name, level=3)
            for step in component.get("steps", []):
                _add_bullet(doc, step)
        doc.add_paragraph()
    else:
        doc.add_paragraph("Interface process not specified.")
        doc.add_paragraph()

    # Optional: embed generated SQL
    if sql.strip():
        _add_heading(doc, "8.3  Generated SQL", level=2)
        sql_para = doc.add_paragraph()
        sql_run = sql_para.add_run(sql)
        sql_run.font.name = "Courier New"
        sql_run.font.size = Pt(8)
        doc.add_paragraph()

    # ── 9. Data Dictionary ────────────────────────────────────────────────────
    _add_heading(doc, "9. Data Dictionary", level=1)
    data_dict = content.get("data_dictionary", [])
    if data_dict:
        _add_table_with_header(
            doc,
            headers=["Dremio Field", "Source Field", "Logic / Derivation"],
            rows=[
                [
                    dd.get("field", ""),
                    dd.get("source_field", ""),
                    dd.get("logic", ""),
                ]
                for dd in data_dict
            ],
        )
    else:
        doc.add_paragraph("Data dictionary not available.")
        doc.add_paragraph()

    # ── 10. Milestones ────────────────────────────────────────────────────────
    _add_heading(doc, "10. Milestones", level=1)
    milestones = content.get("milestones", [])
    if milestones:
        _add_table_with_header(
            doc,
            headers=["Milestone", "Days"],
            rows=[
                [m.get("milestone", ""), str(m.get("days", ""))]
                for m in milestones
            ],
        )
    else:
        doc.add_paragraph("Milestone schedule not specified.")
        doc.add_paragraph()

    # ── 11. Reference ─────────────────────────────────────────────────────────
    _add_heading(doc, "11. Reference", level=1)

    # VDS Source Tables
    _add_heading(doc, "11.1  VDS Source Tables", level=2)
    vds_sources = content.get("vds_source_tables", [])
    if vds_sources:
        _add_table_with_header(
            doc,
            headers=["Table Name", "Full Path"],
            rows=[[vs.get("name", ""), vs.get("path", "")] for vs in vds_sources],
        )
    else:
        doc.add_paragraph("No VDS source tables specified.")
        doc.add_paragraph()

    # VDS Output Paths
    _add_heading(doc, "11.2  VDS Output Paths", level=2)
    vds_outputs = content.get("vds_output_paths", [])
    if vds_outputs:
        _add_table_with_header(
            doc,
            headers=["VDS Name", "Full Dremio Path"],
            rows=[[vo.get("name", ""), vo.get("path", "")] for vo in vds_outputs],
        )
    else:
        if vds_path.strip():
            _add_key_value(doc, "VDS Path", vds_path)
        else:
            doc.add_paragraph("No VDS output paths specified.")
        doc.add_paragraph()

    # Key Business Rules
    _add_heading(doc, "11.3  Key Business Rules", level=2)
    business_rules = content.get("key_business_rules", [])
    if business_rules:
        for rule in business_rules:
            _add_bullet(doc, rule)
    else:
        doc.add_paragraph("No explicit business rules recorded.")
    doc.add_paragraph()

    # Access Control
    _add_heading(doc, "11.4  Access Control", level=2)
    access_control = content.get(
        "access_control",
        "Developer and Tester roles have SELECT access. Write access is restricted to admin.",
    )
    doc.add_paragraph(access_control).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── 12. Sign-off ──────────────────────────────────────────────────────────
    _add_heading(doc, "12. Sign-off", level=1)
    signoff = content.get("signoff", [])
    if signoff:
        _add_table_with_header(
            doc,
            headers=["Role", "Name", "Date"],
            rows=[
                [
                    s.get("role", ""),
                    s.get("name", ""),
                    s.get("date", ""),
                ]
                for s in signoff
            ],
        )
    else:
        doc.add_paragraph("Sign-off table not generated.")

    return doc


# ── Main agent class ─────────────────────────────────────────────────────────────

class DocumentationAgent:
    """Generates a Technical Implementation Document (.docx) from ticket requirements.

    Uses Claude in a single structured JSON call to fill in all document
    sections, then assembles the Word file with python-docx.

    Usage::

        from adl_automated_delivery_pipeline.agents.doc_agent import DocumentationAgent
        from adl_automated_delivery_pipeline.workflows.adl_automated_delivery_pipeline import TicketRequirements

        agent = DocumentationAgent()
        path = agent.generate(reqs, sql=sql_text, vds_path="dremio-db.folder.vds_name")
        print(f"Document saved: {path}")
    """

    def __init__(self) -> None:
        if not settings.ANTHROPIC_API_KEY:
            raise RuntimeError(
                "ANTHROPIC_API_KEY is not set. "
                "Ensure config.env is present and contains ANTHROPIC_API_KEY."
            )
        # Low temperature for precise, well-formed JSON output.
        self._llm = make_claude(model=settings.CLAUDE_MODEL, temperature=0.1)
        logger.info("DocumentationAgent initialised with model=%s", settings.CLAUDE_MODEL)

    # ── Public API ─────────────────────────────────────────────────────────────

    def generate(
        self,
        reqs: Any,
        sql: str = "",
        vds_path: str = "",
    ) -> Path:
        """Generate a .docx Technical Implementation Document.

        Makes a single LLM call to produce all section content as structured
        JSON, then assembles and saves the Word document.

        Args:
            reqs: A TicketRequirements dataclass instance from workflow.py.
            sql: Optional generated SQL query to embed in the document.
            vds_path: Optional Dremio VDS output path (e.g. dremio-db.folder.name).

        Returns:
            Path to the saved .docx file.

        Raises:
            RuntimeError: If the LLM response cannot be parsed as valid JSON.
            OSError: If the output directory cannot be created or file cannot be written.
        """
        logger.info("Generating documentation for ticket %s", reqs.ticket_id)

        content = self._generate_content(reqs, sql, vds_path)
        doc = _assemble_document(reqs, content, sql, vds_path)
        output_path = self._save_document(doc, reqs.ticket_id)

        logger.info("Documentation saved: %s", output_path)
        return output_path

    # ── Private helpers ────────────────────────────────────────────────────────

    def _generate_content(
        self,
        reqs: Any,
        sql: str,
        vds_path: str,
    ) -> dict[str, Any]:
        """Call Claude once to produce all document section content as JSON.

        Args:
            reqs: TicketRequirements dataclass instance.
            sql: Optional SQL string.
            vds_path: Optional VDS path string.

        Returns:
            Parsed dict with all document sections.

        Raises:
            RuntimeError: If the LLM response is not valid JSON after cleanup.
        """
        prompt = _build_doc_prompt(reqs, sql, vds_path)
        logger.info("Calling Claude for document content generation ...")

        try:
            response = self._llm.invoke([
                SystemMessage(
                    content=(
                        "You are a precise technical documentation writer for ASL Airlines. "
                        "Return ONLY a single valid JSON object. "
                        "No markdown code fences. No preamble. No explanation after the JSON."
                    )
                ),
                HumanMessage(content=prompt),
            ])
        except Exception as exc:
            logger.exception("LLM call failed: %s", exc)
            raise RuntimeError(f"LLM call failed: {exc}") from exc

        raw = str(cast(Any, response.content)).strip()

        # Strip markdown fences if the model wraps in ```json ... ```
        if raw.startswith("```"):
            lines = raw.splitlines()
            inner_lines = [
                ln for ln in lines
                if not ln.strip().startswith("```")
            ]
            # Drop leading "json" token if present
            if inner_lines and inner_lines[0].strip().lower() == "json":
                inner_lines = inner_lines[1:]
            raw = "\n".join(inner_lines).strip()

        try:
            content: dict[str, Any] = json.loads(raw)
        except json.JSONDecodeError as exc:
            logger.error("JSON parse error from LLM response: %s", exc)
            logger.debug("Raw LLM response (first 1000 chars): %s", raw[:1000])
            raise RuntimeError(
                f"LLM returned invalid JSON: {exc}. "
                "Check logs for the raw response."
            ) from exc

        logger.info("Document content generated successfully — %d top-level keys", len(content))
        return content

    def _save_document(self, doc: Any, ticket_id: str) -> Path:
        """Save the Document to the output directory.

        Args:
            doc: Fully assembled python-docx Document.
            ticket_id: Jira ticket ID used for the filename.

        Returns:
            Absolute Path to the saved file.

        Raises:
            OSError: If the file cannot be written.
        """
        _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        filename = f"{ticket_id}_{ts}.docx"
        output_path = _OUTPUT_DIR / filename

        try:
            doc.save(str(output_path))
        except OSError as exc:
            logger.exception("Failed to save document to %s: %s", output_path, exc)
            raise

        return output_path


# ── CLI entry point ──────────────────────────────────────────────────────────────

def _demo() -> None:
    """Quick smoke-test: generate a demo document with dummy requirements."""
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(name)s - %(message)s",
    )

    # Import TicketRequirements from the canonical location
    from adl_automated_delivery_pipeline.workflows.adl_automated_delivery_pipeline import TicketRequirements

    demo_reqs = TicketRequirements(
        ticket_id="ADL-1234",
        summary="AMOS MM Discrepancy Report — Rotables Tracking",
        business_requirement=(
            "Generate a Dremio Virtual Dataset that identifies discrepancies between "
            "AMOS maintenance records and rotable part movements, enabling the ASL MRO "
            "team to track un-reconciled transactions by aircraft, part number, and period."
        ),
        source_database="amos_postgres",
        source_tables=["amos.rotables", "amos.rotables_trend", "amos.ac_reg"],
        output_fields=[
            {"name": "ac_reg", "description": "Aircraft registration"},
            {"name": "partno", "description": "APU Part Number"},
            {"name": "serial_no", "description": "Part serial number"},
            {"name": "discrepancy_flag", "description": "Y/N flag for unreconciled transactions"},
            {"name": "last_movement_date", "description": "Date of last recorded movement"},
        ],
        transformations=[
            "Cast movement_date from VARCHAR to DATE using TO_DATE",
            "CASE WHEN reconciled_flag IS NULL THEN 'Y' ELSE 'N' END AS discrepancy_flag",
            "NULLIF(serial_no, '') to handle blank serial numbers",
        ],
        filter_conditions=[
            "movement_date >= DATE '2024-01-01'",
            "ac_type IN ('B737', 'A320')",
        ],
        acceptance_criteria=[
            "All rotable movements since Jan 2024 appear in the output",
            "Discrepancy flag is Y for any record missing a reconciliation entry",
            "VDS returns results within 5 seconds on a 100k-row dataset",
        ],
        extra_notes="Priority: High. Stakeholder sign-off required before UAT.",
        raw_description="Full Jira description would appear here.",
    )

    demo_sql = """-- VDS: adl_mm_discrepancy
-- Ticket: ADL-1234 | Generated by DnT Dremio Agent
SELECT
    r.ac_reg,
    r.partno,
    NULLIF(r.serial_no, '') AS serial_no,
    TO_DATE(r.movement_date, 'YYYY-MM-DD') AS last_movement_date,
    CASE WHEN r.reconciled_flag IS NULL THEN 'Y' ELSE 'N' END AS discrepancy_flag
FROM amos_postgres.amos.rotables r
JOIN amos_postgres.amos.ac_reg a ON a.ac_reg = r.ac_reg
WHERE TO_DATE(r.movement_date, 'YYYY-MM-DD') >= DATE '2024-01-01'
  AND a.ac_type IN ('B737', 'A320')"""

    agent = DocumentationAgent()
    path = agent.generate(
        reqs=demo_reqs,
        sql=demo_sql,
        vds_path="dremio-db.amos_maintenance.adl_mm_discrepancy",
    )
    print(f"\nDocument generated: {path}\n")


if __name__ == "__main__":
    _demo()


