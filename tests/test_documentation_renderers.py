"""Unit tests for brand helpers and renderers (assert on the docx object model)."""

from __future__ import annotations

import pytest
from docx import Document

from adl_automated_delivery_pipeline.documentation import brand


@pytest.mark.unit
def test_add_heading_adds_text() -> None:
    doc = Document()
    brand.add_heading(doc, "Section", level=1)
    para = doc.paragraphs[-1]
    assert para.runs[0].text == "Section"
    # Colour is inherited from OCC template styles — not overridden inline.


@pytest.mark.unit
def test_add_table_builds_header_plus_rows_and_pads_short_rows() -> None:
    doc = Document()
    brand.add_table(doc, ["A", "B"], [["1", "2"], ["3"]])
    table = doc.tables[-1]
    assert len(table.rows) == 3            # header + 2 data rows
    assert table.rows[0].cells[0].text == "A"
    assert table.rows[2].cells[1].text == ""  # short row padded


from pathlib import Path

from adl_automated_delivery_pipeline.documentation.context import DocContext
from adl_automated_delivery_pipeline.documentation.renderers import (
    base as renderers_base,
)
from adl_automated_delivery_pipeline.documentation import renderers  # noqa: F401  (registers md/docx)


@pytest.mark.unit
def test_get_renderer_unknown_format_lists_registered() -> None:
    with pytest.raises(ValueError) as exc:
        renderers_base.get_renderer("xml")
    assert "md" in str(exc.value)


@pytest.mark.unit
def test_markdown_renderer_writes_file_verbatim(tmp_path: Path) -> None:
    r = renderers_base.get_renderer("md")
    out = r.render("# Hello\n\nBody", tmp_path / "doc.md", DocContext(title="x"))
    assert out.read_text(encoding="utf-8") == "# Hello\n\nBody"
    assert out.suffix == ".md"


from docx import Document as _Doc

_MD = """# Title

Intro paragraph.

## Risks

| Risk | Mitigation |
|------|------------|
| Schema drift | Validate with EXPLAIN |
| Large scan | Add filter |

## Steps

- First step
- Second step

```sql
SELECT 1
```
"""


@pytest.mark.unit
def test_docx_renderer_maps_markdown_to_docx_model(tmp_path: Path) -> None:
    r = renderers_base.get_renderer("docx")
    out = r.render(_MD, tmp_path / "doc.docx", DocContext(title="Title"))
    assert out.suffix == ".docx"

    doc = _Doc(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Title" in texts
    assert "Intro paragraph." in texts
    assert any("First step" in t for t in texts)  # may be prefixed with bullet char
    # one table with header + 2 rows
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "Risk"
    assert doc.tables[0].rows[1].cells[1].text == "Validate with EXPLAIN"
    assert len(doc.tables[0].rows) == 3


@pytest.mark.unit
def test_docx_renderer_nested_bullets_all_stay_bullets(tmp_path: Path) -> None:
    r = renderers_base.get_renderer("docx")
    out = r.render("- Before\n  - Nested\n- After\n", tmp_path / "n.docx", DocContext(title="t"))
    doc = _Doc(str(out))
    _BULLET_STYLES = {"List Bullet", "List Paragraph"}
    bullets = [p.text.lstrip("\xa0• ") for p in doc.paragraphs if p.style.name in _BULLET_STYLES]
    assert "Before" in bullets
    assert "Nested" in bullets
    assert "After" in bullets  # before the fix this rendered as a Normal paragraph
